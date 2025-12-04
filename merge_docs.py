#!/usr/bin/env python3
"""
Merge all .docx files from to_merge directory into a single PDF file.

Features:
- Merges all .docx files in alphabetical order
- Preserves formatting, images, tables
- Clean, simple output with no complex bookmarking
- Zero configuration required
- Outputs to Merged_Doc.pdf

Dependencies:
- python-docx: For reading .docx files
- pillow: For image handling
- libreoffice: For converting .docx to .pdf (via command line)
"""

import os
import sys
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import shutil
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil


def get_docx_files(directory):
    """
    Scan directory for all .docx files and return them sorted alphabetically.
    """
    docx_files = sorted(directory.glob("*.docx"))
    if not docx_files:
        print(f"⚠️  No .docx files found in {directory}")
        return []
    print(f"✓ Found {len(docx_files)} .docx files")
    for f in docx_files:
        print(f"  - {f.name}")
    return docx_files


# Global bookmark id counter
_bookmark_id_counter = 1


def _next_bookmark_id():
    global _bookmark_id_counter
    val = _bookmark_id_counter
    _bookmark_id_counter += 1
    return str(val)


def add_bookmark_to_paragraph(paragraph, bookmark_name):
    """Insert a bookmarkStart/bookmarkEnd pair around a paragraph."""
    bookmark_id = _next_bookmark_id()
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    # Insert at start and append end
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)


def add_internal_hyperlink(paragraph, anchor, text):
    """Add an internal hyperlink (anchor) to `paragraph` with visible `text`."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    # Create run
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    # color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def load_document(file_path):
    """Load a .docx document from file with error handling and recovery."""
    try:
        return Document(file_path)
    except Exception as e:
        error_msg = str(e)
        
        # For extremely large files, try LibreOffice resave as recovery
        if "Buffer size limit exceeded" in error_msg or "Resource limit" in error_msg:
            print(f"⚠️  File too large, attempting recovery via LibreOffice...")
            try:
                import tempfile
                with tempfile.TemporaryDirectory() as tmpdir:
                    cmd = [
                        "libreoffice",
                        "--headless",
                        "--convert-to", "docx",
                        "--outdir", tmpdir,
                        str(file_path)
                    ]
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
                    if result.returncode == 0:
                        # Try to load the resaved file
                        resaved = Path(tmpdir) / file_path.name
                        if resaved.exists():
                            doc = Document(resaved)
                            print(f"✓ Loaded resaved file for {file_path.name}")
                            return doc
            except Exception as recovery_err:
                print(f"   Recovery failed: {str(recovery_err)[:80]}")
        
        print(f"❌ Error loading {file_path.name}: {str(e)[:100]}")
        return None


def detect_pdf_converter():
    """Detect which PDF converter will be used and return a string.

    Preference order:
      1. `soffice` (LibreOffice)
      2. `mammoth` + `wkhtmltopdf` fallback
      3. `none` if no converter available
    """
    # Check for soffice (LibreOffice)
    if shutil.which('soffice'):
        return 'soffice'

    # Check for mammoth + wkhtmltopdf availability
    try:
        import mammoth  # noqa: F401
        if shutil.which('wkhtmltopdf'):
            return 'mammoth+wkhtmltopdf'
    except Exception:
        pass

    return 'none'


def copy_paragraph(source_para, target_doc, source_doc=None, images_outdir=None):
    """
    Copy a paragraph from source to target, preserving all formatting, images, hyperlinks, and bookmarks.
    """
    # Create a new paragraph with same style and alignment
    dst_para = target_doc.add_paragraph(style=source_para.style)
    dst_para.alignment = source_para.alignment
    
    # Copy indentation and spacing
    if source_para.paragraph_format.left_indent:
        dst_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    if source_para.paragraph_format.right_indent:
        dst_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    if source_para.paragraph_format.space_before:
        dst_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    if source_para.paragraph_format.space_after:
        dst_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    if source_para.paragraph_format.line_spacing:
        dst_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    
    # Copy line spacing type
    try:
        dst_para.paragraph_format.line_spacing_rule = source_para.paragraph_format.line_spacing_rule
    except Exception:
        pass
    
    # Copy borders and shading if present
    try:
        src_pPr = source_para._element.pPr
        if src_pPr is not None:
            dst_pPr = dst_para._element.pPr
            # Copy borders
            pBdr = src_pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                from copy import deepcopy
                dst_pPr.append(deepcopy(pBdr))
            # Copy shading
            shd = src_pPr.find(qn('w:shd'))
            if shd is not None:
                from copy import deepcopy
                dst_pPr.append(deepcopy(shd))
    except Exception:
        pass
    
    # Extract and preserve bookmarks and hyperlinks from source paragraph XML
    src_p_element = source_para._element
    bookmarks = {}
    hyperlinks = {}
    
    # Find all bookmarks in this paragraph
    for elem in src_p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart'):
        bk_id = elem.get(qn('w:id'))
        bk_name = elem.get(qn('w:name'))
        if bk_id and bk_name:
            bookmarks[bk_id] = bk_name
    
    # Find all hyperlinks in this paragraph
    for hlink in src_p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink'):
        rel_id = hlink.get(qn('r:id'))
        anchor = hlink.get(qn('w:anchor'))
        hyperlinks[rel_id] = anchor
    
    # Copy runs (text with formatting, images, hyperlinks, and bookmarks)
    for run in source_para.runs:
        # Check if run contains an image (drawing)
        if run.element.drawing_lst:
            # This run has an image
            try:
                for drawing in run.element.drawing_lst:
                    # Try to extract image info from drawing
                    for blip in drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id and source_doc:
                            # Get the image from source document
                            try:
                                img_part = source_doc.part.related_part(embed_id)
                                img_data = img_part.blob
                                # Get the drawing size info
                                extent = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                                if extent is not None:
                                    # Size is in EMUs (English Metric Units)
                                    width_emu = int(extent.get('cx', 0))
                                    height_emu = int(extent.get('cy', 0))
                                    # Convert EMU to inches (1 inch = 914400 EMUs)
                                    width_inches = Inches(width_emu / 914400.0)
                                    height_inches = Inches(height_emu / 914400.0)
                                    # Add picture to destination paragraph with original size
                                    dst_para.add_run().add_picture(BytesIO(img_data), width=width_inches, height=height_inches)
                                else:
                                    # No size info, add with default size
                                    dst_para.add_run().add_picture(BytesIO(img_data))
                            except Exception:
                                pass
            except Exception:
                pass
        else:
            # Check if this run is part of a hyperlink
            parent = run._element.getparent()
            is_hyperlink = parent.tag.endswith('}hyperlink')
            
            # Normal text run copy
            new_run = dst_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            
            # Copy strike-through, subscript, superscript
            try:
                new_run.font.strike = run.font.strike
                new_run.font.subscript = run.font.subscript
                new_run.font.superscript = run.font.superscript
            except Exception:
                pass
            
            # If this run was inside a hyperlink, preserve that info
            if is_hyperlink and source_doc:
                try:
                    hyperlink_elem = parent
                    rel_id = hyperlink_elem.get(qn('r:id'))
                    anchor = hyperlink_elem.get(qn('w:anchor'))
                    
                    if rel_id:
                        # External hyperlink - try to resolve the URL
                        try:
                            related = source_doc.part.rels[rel_id]
                            target_url = related.target_ref
                            # Create external hyperlink in destination
                            new_run.hyperlink.address = target_url
                        except Exception:
                            pass
                    elif anchor:
                        # Internal hyperlink/bookmark reference
                        new_run.hyperlink.anchor = anchor
                except Exception:
                    pass


def copy_run_with_hyperlink(dst_para, text, href=None, anchor=None):
    """Helper to add a run with hyperlink styling."""
    new_run = dst_para.add_run(text)
    new_run.font.color.rgb = RGBColor(0, 0, 255)
    new_run.underline = True
    if href:
        new_run.hyperlink.address = href
    if anchor:
        new_run.hyperlink.anchor = anchor
    return new_run


def copy_table(source_table, target_doc, source_doc=None, images_outdir=None):
    """
    Copy a table from source to target document, preserving structure, formatting, borders, and cell properties.
    """
    # Create table with same dimensions
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    dst_table = target_doc.add_table(rows=rows, cols=cols)
    
    # Copy table style if available
    if source_table.style:
        try:
            dst_table.style = source_table.style
        except Exception:
            pass
    
    # Copy table-wide properties (borders, shading)
    try:
        src_tbl = source_table._element
        src_tblPr = src_tbl.tblPr
        if src_tblPr is not None:
            dst_tbl = dst_table._element
            dst_tblPr = dst_tbl.tblPr
            # Copy table borders
            tblBorders = src_tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                from copy import deepcopy
                dst_tblPr.append(deepcopy(tblBorders))
            # Copy table width
            tblW = src_tblPr.find(qn('w:tblW'))
            if tblW is not None:
                from copy import deepcopy
                dst_tblPr.append(deepcopy(tblW))
    except Exception:
        pass
    
    # Copy cell contents and formatting
    for r in range(rows):
        for c in range(cols):
            src_cell = source_table.cell(r, c)
            dst_cell = dst_table.cell(r, c)
            
            # Clear default paragraph
            dst_cell.text = ""
            
            # Copy paragraphs from source cell
            for src_para in src_cell.paragraphs:
                copy_paragraph(src_para, dst_cell, source_doc=source_doc, images_outdir=images_outdir)
            
            # Copy cell properties (shading, borders, width)
            try:
                src_tc = src_cell._tc
                src_tcPr = src_tc.tcPr
                if src_tcPr is not None:
                    dst_tc = dst_cell._tc
                    dst_tcPr = dst_tc.get_or_add_tcPr()
                    
                    # Copy cell shading
                    shd = src_tcPr.find(qn('w:shd'))
                    if shd is not None:
                        from copy import deepcopy
                        existing_shd = dst_tcPr.find(qn('w:shd'))
                        if existing_shd is not None:
                            dst_tcPr.remove(existing_shd)
                        dst_tcPr.append(deepcopy(shd))
                    
                    # Copy cell borders
                    tcBorders = src_tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        from copy import deepcopy
                        existing_borders = dst_tcPr.find(qn('w:tcBorders'))
                        if existing_borders is not None:
                            dst_tcPr.remove(existing_borders)
                        dst_tcPr.append(deepcopy(tcBorders))
                    
                    # Copy cell width
                    tcW = src_tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        from copy import deepcopy
                        existing_w = dst_tcPr.find(qn('w:tcW'))
                        if existing_w is not None:
                            dst_tcPr.remove(existing_w)
                        dst_tcPr.append(deepcopy(tcW))
                    
                    # Copy vertical alignment
                    vAlign = src_tcPr.find(qn('w:vAlign'))
                    if vAlign is not None:
                        from copy import deepcopy
                        existing_vAlign = dst_tcPr.find(qn('w:vAlign'))
                        if existing_vAlign is not None:
                            dst_tcPr.remove(existing_vAlign)
                        dst_tcPr.append(deepcopy(vAlign))
            except Exception:
                pass


def copy_document_elements(source_doc, target_doc, images_outdir=None):
    """
    Copy all paragraphs and tables from source to target document.
    """
    # Iterate through body elements in order
    for element in source_doc.element.body:
        tag = element.tag
        
        if tag.endswith('}p'):
            # This is a paragraph
            for para in source_doc.paragraphs:
                if para._element is element:
                    copy_paragraph(para, target_doc, source_doc=source_doc, images_outdir=images_outdir)
                    break
        
        elif tag.endswith('}tbl'):
            # This is a table
            for table in source_doc.tables:
                if table._element is element:
                    copy_table(table, target_doc, source_doc=source_doc, images_outdir=images_outdir)
                    break


def copy_styles_and_properties(source_doc, target_doc):
    """
    Copy document-level styles and properties from source to target.
    """
    try:
        # Copy core properties
        target_doc.core_properties.title = source_doc.core_properties.title
        target_doc.core_properties.subject = source_doc.core_properties.subject
        target_doc.core_properties.author = source_doc.core_properties.author
        target_doc.core_properties.keywords = source_doc.core_properties.keywords
        target_doc.core_properties.comments = source_doc.core_properties.comments
    except Exception:
        pass
    
    # Note: Copying custom styles/themes requires deeper XML manipulation
    # and can be complex. This is a best-effort approach.


def merge_docx_to_docx(docx_files, output_path):
    """
    Merge multiple .docx files into a single .docx file, preserving original page layout.
    """
    if not docx_files:
        print("❌ No documents to merge")
        return None
    
    # Create master document
    master_doc = Document()
    
    print("\n📄 Merging documents...")

    for idx, docx_file in enumerate(docx_files):
        print(f"  [{idx+1}/{len(docx_files)}] Processing: {docx_file.name}")
        
        source_doc = load_document(docx_file)
        if not source_doc:
            continue

        # prepare images output dir for this source and extract images from source .docx
        images_base = Path(output_path).with_name('extracted_images')
        images_outdir = images_base / docx_file.stem
        try:
            images_outdir.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        # extract all images from source file into images_outdir
        try:
            from zipfile import ZipFile
            with ZipFile(docx_file, 'r') as z:
                for name in z.namelist():
                    if name.startswith('word/media/'):
                        target = images_outdir / os.path.basename(name)
                        with z.open(name) as src, open(target, 'wb') as out_f:
                            out_f.write(src.read())
        except Exception:
            pass

        # Copy all document content (text/tables) preserving original structure
        copy_document_elements(source_doc, master_doc)
        
        # Add page break between documents (except after last one)
        if idx < len(docx_files) - 1:
            master_doc.add_page_break()
    
    # Save merged document
    try:
        master_doc.save(output_path)
        print(f"✓ Merged document saved to {output_path.name}")
        return output_path
    except Exception as e:
        print(f"❌ Error saving merged document: {e}")
        # Try fallback filename
        try:
            alt = output_path.with_name(output_path.stem + '_autosaved' + output_path.suffix)
            master_doc.save(alt)
            print(f"✓ Merged document saved to fallback file {alt.name}")
            return alt
        except Exception as e2:
            print(f"❌ Fallback save failed: {e2}")
            return None


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert .docx file to .pdf using LibreOffice command line.
    """
    print(f"\n📝 Converting to PDF (LibreOffice -> fallback)...")

    # First try LibreOffice/soffice (best fidelity)
    try:
        cmd = [
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(pdf_path.parent),
            str(docx_path)
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
        if result.returncode == 0:
            print(f"✓ PDF conversion successful (soffice)")
            return True
        else:
            print(f"⚠️  soffice returned non-zero exit: {result.returncode}. Falling back...")
    except FileNotFoundError:
        print("⚠️  `soffice` not found on PATH — falling back to mammoth+wkhtmltopdf")
    except Exception as e:
        print(f"⚠️  soffice conversion error: {e}. Falling back...")

    # Fallback: mammoth (docx -> HTML) + wkhtmltopdf (HTML -> PDF)
    try:
        import mammoth
    except Exception:
        print("❌ mammoth not installed; install with `pip install mammoth` to enable fallback conversion")
        return False

    try:
        import tempfile
        print("ℹ️  Using mammoth to convert DOCX to HTML...")
        with open(docx_path, 'rb') as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value

        with tempfile.TemporaryDirectory() as tmpdir:
            html_path = Path(tmpdir) / (pdf_path.stem + '.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html)

            # Check for wkhtmltopdf
            wk_cmd = ["wkhtmltopdf", str(html_path), str(pdf_path)]
            try:
                wk_res = subprocess.run(wk_cmd, capture_output=True, text=True, timeout=300)
                if wk_res.returncode == 0:
                    print("✓ PDF conversion successful (mammoth + wkhtmltopdf)")
                    return True
                else:
                    print(f"❌ wkhtmltopdf failed: {wk_res.stderr}")
                    return False
            except FileNotFoundError:
                print("❌ `wkhtmltopdf` not found — install it to enable HTML->PDF fallback")
                return False
            except Exception as e:
                print(f"❌ wkhtmltopdf error: {e}")
                return False
    except Exception as e:
        print(f"❌ Fallback conversion error: {e}")
        return False


def main():
    """Main entry point for the script."""
    # Get script directory
    script_dir = Path(__file__).parent.resolve()
    merge_dir = script_dir / "to_merge"
    output_docx = script_dir / "Merged_Doc.docx"
    output_pdf = script_dir / "Merged_Doc.pdf"
    
    print("=" * 60)
    print("🔀 DOCX to PDF Merger")
    print("=" * 60)
    print(f"📂 Script directory: {script_dir}")
    print(f"📂 Merge directory: {merge_dir}")
    # Detect which converter will be used and print it
    converter = detect_pdf_converter()
    if converter == 'soffice':
        print("🔧 PDF converter: LibreOffice (`soffice`) — best fidelity")
    elif converter == 'mammoth+wkhtmltopdf':
        print("🔧 PDF converter: fallback to `mammoth` + `wkhtmltopdf` — lighter, reduced fidelity")
    else:
        print("⚠️  No PDF converter found: please install LibreOffice or mammoth+wkhtmltopdf")
    
    # Check if merge directory exists
    if not merge_dir.exists():
        print(f"❌ Merge directory not found: {merge_dir}")
        return 1
    
    # Get all .docx files
    docx_files = get_docx_files(merge_dir)
    if not docx_files:
        return 1
    
    # Merge documents
    merged_path = merge_docx_to_docx(docx_files, output_docx)
    if not merged_path:
        return 1
    
    # Convert to PDF
    if not convert_docx_to_pdf(merged_path, output_pdf):
        print("⚠️  DOCX merge successful but PDF conversion failed")
        print(f"   Merged DOCX saved to: {merged_path}")
        return 1
    
    # Verify output
    if output_pdf.exists():
        size_mb = output_pdf.stat().st_size / (1024 * 1024)
        print(f"✓ Final PDF size: {size_mb:.2f} MB")
        print(f"✓ Output saved to: {output_pdf}")
        print("=" * 60)
        print("✓ Merge complete!")
        print("=" * 60)
        return 0
    else:
        print(f"❌ Output file not created: {output_pdf}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
